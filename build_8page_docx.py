"""Render the condensed paper_short.md into the skeleton's DOCX styling.

Adapted from build_docx.py. Opens mcp-dos-paper-skeleton.docx so all of its named
styles (Title, Heading 1-4, List Paragraph, Normal, theme fonts/colors) are inherited,
clears the template body, and rewrites the short paper content with the skeleton's
signature table styling (blue header rows, light-grey borders, alternating row fill).

Adds image embedding: a markdown image directive on its own line, of the form
    ![Figure N. caption text](experiments/figures/figX.png)
is rendered as the embedded PNG (scaled to ~6.0in text width, aspect preserved,
centered) followed by a centered italic caption paragraph.

Outputs mcp-dos-paper-SHORT.docx. Does not touch mcp-dos-paper-FILLED.docx or the skeleton.
"""
import os
import re
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "paper_8page.md")
SKELETON = os.path.join(HERE, "mcp-security-paper.docx")
OUT = os.path.join(HERE, "mcp-security-paper.docx")

BLUE = "0B4F6C"      # header fill
WHITE = "FFFFFF"
ALT = "EEF3F5"       # alternating body row
BORDER = "BFC7CC"    # light grey borders
TWO_COL = False      # single-column layout: figures/tables render full width, no section breaks
IMG_WIDTH_IN = 5.2   # target image width (near full text width so charts are clearly legible)
BODY_PT = 10.0       # comfortable single-column body font
TABLE_PT = 8.0       # compact table cell font (headers + body)
HEADING_STYLES = {"Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"}

doc = Document(SKELETON)

STYLES = {}
for s in doc.styles:
    if s.name not in STYLES:
        STYLES[s.name] = s

# --- clear template body but keep section properties (page size/margins) ---
body = doc.element.body
sectPr = body.find(qn("w:sectPr"))
for child in list(body):
    if child is not sectPr:
        body.remove(child)

# --- two-column layout with full-width section breaks around figures --------
def _set_cols(sp, num):
    """Set (replace) the w:cols child of a sectPr to `num` columns."""
    existing = sp.find(qn("w:cols"))
    if existing is not None:
        sp.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "425")  # ~0.3in gutter, twips
    sp.append(cols)

# single-column body (TWO_COL=False)
_set_cols(sectPr, 1)

# --- density: shrink margins, Normal + heading style sizes, tighten spacing ---
def _set_margins(sp, inch=0.75):
    pgMar = sp.find(qn("w:pgMar"))
    if pgMar is not None:
        tw = str(int(inch * 1440))
        for side in ("top", "bottom", "left", "right"):
            pgMar.set(qn("w:" + side), tw)

_set_margins(sectPr, 0.75)  # cloned into every later section break via _clone_page_geometry

def _style_size(name, pt):
    if name in STYLES:
        try:
            STYLES[name].font.size = Pt(pt)
        except Exception:
            pass

_style_size("Normal", BODY_PT)
_style_size("Heading 1", 11)
_style_size("Heading 2", 10)
_style_size("Heading 3", 9.5)
_style_size("Title", 15)
try:
    npf = STYLES["Normal"].paragraph_format
    npf.space_after = Pt(2); npf.space_before = Pt(0); npf.line_spacing = 1.0
except Exception:
    pass
for _hs in ("Heading 1", "Heading 2", "Heading 3"):
    if _hs in STYLES:
        try:
            hpf = STYLES[_hs].paragraph_format
            hpf.space_before = Pt(4); hpf.space_after = Pt(2)
        except Exception:
            pass

def _clone_page_geometry(target_sp):
    """Copy pgSz/pgMar from the body sectPr so a break keeps the same page setup."""
    for tag in ("w:pgSz", "w:pgMar"):
        src = sectPr.find(qn(tag))
        if src is not None:
            import copy
            target_sp.append(copy.deepcopy(src))

def end_section_on(paragraph, num_cols):
    """Terminate a continuous section (with num_cols columns) at this paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    # remove any pre-existing sectPr on this paragraph
    old = pPr.find(qn("w:sectPr"))
    if old is not None:
        pPr.remove(old)
    sp = OxmlElement("w:sectPr")
    stype = OxmlElement("w:type")
    stype.set(qn("w:val"), "continuous")
    sp.append(stype)
    _clone_page_geometry(sp)
    _set_cols(sp, num_cols)
    pPr.append(sp)

# --- helpers -----------------------------------------------------------------
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def table_borders(table, color=BORDER, sz="4"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)

def set_table_cell_margins(table, top=14, bottom=14, left=43, right=43):
    """Set uniform, tight cell margins (twips) on all cells via tblCellMar.
    43 twips ~= 0.03in left/right, 14 twips ~= 0.01in top/bottom."""
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)

def set_repeat_header(row):
    """Mark a table row to repeat as a header on every page it spans."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)

def _tighten_cell(cell):
    """Zero out paragraph spacing inside a table cell so rows stay short."""
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

INLINE = re.compile(r"(==.+?==|\*\*.+?\*\*|~~.+?~~|`.+?`|\*.+?\*)")
IMAGE = re.compile(r"^!\[(?P<caption>.*)\]\((?P<path>[^)]+)\)\s*$")

def _emit(p, text, bold, italic, strike, code, color, highlight=False):
    if not text:
        return
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if strike:
        r.font.strike = True
    if code:
        r.font.name = "Consolas"; r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if highlight:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW

def add_runs(p, text, base_bold=False, base_color=None, base_strike=False, base_highlight=False):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _emit(p, text[pos:m.start()], base_bold, False, base_strike, False, base_color, base_highlight)
        tok = m.group(0)
        if tok.startswith("=="):
            add_runs(p, tok[2:-2], base_bold=base_bold, base_color=base_color, base_strike=base_strike, base_highlight=True)
        elif tok.startswith("**"):
            add_runs(p, tok[2:-2], base_bold=True, base_color=base_color, base_strike=base_strike, base_highlight=base_highlight)
        elif tok.startswith("~~"):
            add_runs(p, tok[2:-2], base_bold=base_bold, base_color=base_color, base_strike=True, base_highlight=base_highlight)
        elif tok.startswith("`"):
            _emit(p, tok[1:-1], False, False, base_strike, True, None, base_highlight)
        elif tok.startswith("*"):
            _emit(p, tok[1:-1], base_bold, True, base_strike, False, base_color, base_highlight)
        pos = m.end()
    if pos < len(text):
        _emit(p, text[pos:], base_bold, False, base_strike, False, base_color, base_highlight)

def para(text, style=None, bold=False, color=None, size=None, align=None, space_after=2):
    p = doc.add_paragraph()
    if style and style in STYLES:
        p.style = STYLES[style]
    add_runs(p, text, base_bold=bold, base_color=color)
    if size:
        for r in p.runs:
            r.font.size = Pt(size)
    elif style not in HEADING_STYLES:
        # explicit body size — style inheritance alone did not override the skeleton
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(BODY_PT)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_image(caption, rel_path):
    """Embed a PNG full-page-width in a 1-column continuous section, centered,
    with a caption below. The surrounding body remains two-column."""
    abs_path = rel_path if os.path.isabs(rel_path) else os.path.join(HERE, rel_path.replace("/", os.sep))
    with Image.open(abs_path) as im:
        w_px, h_px = im.size
    # In two-column mode, isolate the figure in its own full-width section.
    if TWO_COL and doc.paragraphs:
        end_section_on(doc.paragraphs[-1], 2)
    width_in = IMG_WIDTH_IN
    # preserve aspect ratio; python-docx sets height automatically when only width given
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(abs_path, width=Inches(width_in))
    pic_p.paragraph_format.space_after = Pt(3)
    # caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(cap, "*" + caption + "*")
    for r in cap.runs:
        r.font.size = Pt(8.5)
    cap.paragraph_format.space_after = Pt(6)
    # Close the figure's full-width section; body resumes two-column after.
    if TWO_COL:
        end_section_on(cap, 1)
    return abs_path, (w_px, h_px)

def render_table(rows):
    header, data = rows[0], rows[1:]
    # In two-column mode, wide tables get their own full-width section so they
    # are not crushed into a ~3.2in strip. In single-column mode this is unneeded.
    full_width = TWO_COL and len(header) >= 5
    if full_width and doc.paragraphs:
        end_section_on(doc.paragraphs[-1], 2)
    t = doc.add_table(rows=1, cols=len(header))
    t.autofit = True
    # make the table span the full text width so wide tables stay readable
    tblPr = t._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct"); tblW.set(qn("w:w"), "5000")  # 100%
    table_borders(t)
    set_table_cell_margins(t)
    set_repeat_header(t.rows[0])
    for j, cell_text in enumerate(header):
        c = t.rows[0].cells[j]
        c.paragraphs[0].text = ""
        add_runs(c.paragraphs[0], cell_text, base_bold=True, base_color=WHITE)
        for r in c.paragraphs[0].runs:
            r.font.size = Pt(TABLE_PT)
        shade_cell(c, BLUE)
        _tighten_cell(c)
    for i, row in enumerate(data):
        cells = t.add_row().cells
        fill = ALT if i % 2 == 1 else WHITE
        for j in range(len(header)):
            txt = row[j] if j < len(row) else ""
            c = cells[j]
            c.paragraphs[0].text = ""
            add_runs(c.paragraphs[0], txt)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(TABLE_PT)
            shade_cell(c, fill)
            _tighten_cell(c)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    if full_width:
        end_section_on(spacer, 1)
    return t

def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]

# --- parse markdown ----------------------------------------------------------
with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
first_heading = True
n = len(lines)
images_embedded = []
while i < n:
    line = lines[i]
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    if stripped == "---":
        i += 1
        continue

    # image directive (own line)
    im = IMAGE.match(stripped)
    if im:
        p, dims = add_image(im.group("caption"), im.group("path"))
        images_embedded.append(p)
        i += 1
        continue

    # code fence
    if stripped.startswith("```"):
        code = []
        i += 1
        while i < n and not lines[i].strip().startswith("```"):
            code.append(lines[i])
            i += 1
        i += 1
        p = doc.add_paragraph()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:val"), "clear"); shade.set(qn("w:color"), "auto"); shade.set(qn("w:fill"), "F2F2F2")
        p.paragraph_format.element.get_or_add_pPr().append(shade)
        for k, cl in enumerate(code):
            r = p.add_run(cl)
            r.font.name = "Consolas"; r.font.size = Pt(9)
            if k < len(code) - 1:
                r.add_break()
        continue

    # table block
    if stripped.startswith("|"):
        block = []
        while i < n and lines[i].strip().startswith("|"):
            block.append(lines[i])
            i += 1
        rows = [split_row(b) for b in block]
        rows = [r for r in rows if not all(set(c) <= set("-: ") and c for c in r)]
        if rows:
            render_table(rows)
        continue

    # heading
    if stripped.startswith("#"):
        m = re.match(r"(#+)\s+(.*)", stripped)
        level = len(m.group(1))
        text = m.group(2)
        if level == 1 and first_heading:
            para(text, style="Title")
            para("Giorgi Iashvili and Maksim Iavich",
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
            para("School of Technology, Caucasus University, Tbilisi, Georgia",
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
            para("Corresponding author: giiashvili@cu.edu.ge",
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
            first_heading = False
        else:
            style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 3")
            para(text, style=style)
        i += 1
        continue

    # bullet
    if stripped.startswith("- "):
        p = doc.add_paragraph()
        if "List Paragraph" in STYLES:
            p.style = STYLES["List Paragraph"]
        add_runs(p, "•  " + stripped[2:])
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(BODY_PT)
        p.paragraph_format.space_after = Pt(2)
        i += 1
        continue

    para(stripped)
    i += 1

doc.save(OUT)
print("saved", OUT)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
print("images embedded:", len(images_embedded))
for p in images_embedded:
    print("  -", p)
